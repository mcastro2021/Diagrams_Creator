import os
import re
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
import json

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.xlsx': self._process_excel,
            '.csv': self._process_csv,
            '.txt': self._process_text,
            '.json': self._process_json
        }
    
    def process_document(self, filepath):
        """Procesa un documento y extrae su contenido estructurado"""
        try:
            file_extension = os.path.splitext(filepath)[1].lower()
            
            if file_extension in self.supported_extensions:
                return self.supported_extensions[file_extension](filepath)
            else:
                return {
                    'type': 'error',
                    'message': f'Extensión de archivo no soportada: {file_extension}'
                }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando documento: {str(e)}'
            }
    
    def _process_pdf(self, filepath):
        """Procesa archivos PDF"""
        try:
            reader = PdfReader(filepath)
            text_content = ""
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
            
            # Detectar información de redes
            network_info = self._extract_network_info(text_content)
            
            return {
                'type': 'network' if network_info else 'text',
                'content': text_content,
                'pages': len(reader.pages),
                'format': 'pdf',
                'network_info': network_info
            }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando PDF: {str(e)}'
            }
    
    def _process_docx(self, filepath):
        """Procesa archivos Word"""
        try:
            doc = Document(filepath)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Detectar información de redes
            network_info = self._extract_network_info(text_content)
            
            return {
                'type': 'network' if network_info else 'text',
                'content': text_content,
                'paragraphs': len(doc.paragraphs),
                'format': 'docx',
                'network_info': network_info
            }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando Word: {str(e)}'
            }
    
    def _process_excel(self, filepath):
        """Procesa archivos Excel"""
        try:
            workbook = load_workbook(filepath, data_only=True)
            sheet = workbook.active
            
            # Extraer datos de la hoja
            data = []
            headers = []
            
            for row in sheet.iter_rows(values_only=True):
                if not headers:
                    headers = [str(cell) if cell else f'Columna_{i+1}' for i, cell in enumerate(row)]
                else:
                    row_data = {}
                    for i, cell in enumerate(row):
                        if i < len(headers):
                            row_data[headers[i]] = str(cell) if cell else ''
                    data.append(row_data)
            
            # Detectar si es información de redes
            network_info = self._detect_network_excel(headers, data)
            
            return {
                'type': 'network' if network_info else 'spreadsheet',
                'content': {'headers': headers, 'data': data},
                'format': 'excel',
                'network_info': network_info
            }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando Excel: {str(e)}'
            }
    
    def _process_csv(self, filepath):
        """Procesa archivos CSV"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                lines = file.readlines()
            
            if not lines:
                return {
                    'type': 'error',
                    'message': 'Archivo CSV vacío'
                }
            
            # Parsear CSV
            headers = lines[0].strip().split(',')
            data = []
            
            for line in lines[1:]:
                values = line.strip().split(',')
                row_data = {}
                for i, value in enumerate(values):
                    if i < len(headers):
                        row_data[headers[i]] = value
                data.append(row_data)
            
            # Detectar si es información de redes
            network_info = self._detect_network_csv(headers, data)
            
            return {
                'type': 'network' if network_info else 'spreadsheet',
                'content': {'headers': headers, 'data': data},
                'format': 'csv',
                'network_info': network_info
            }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando CSV: {str(e)}'
            }
    
    def _process_text(self, filepath):
        """Procesa archivos de texto"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Detectar información de redes
            network_info = self._extract_network_info(text_content)
            
            return {
                'type': 'network' if network_info else 'text',
                'content': text_content,
                'characters': len(text_content),
                'format': 'text',
                'network_info': network_info
            }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando texto: {str(e)}'
            }
    
    def _process_json(self, filepath):
        """Procesa archivos JSON"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                json_content = json.loads(file.read())
            
            # Detectar si es información de redes
            network_info = self._detect_network_json(json_content)
            
            return {
                'type': 'network' if network_info else 'json',
                'content': json_content,
                'format': 'json',
                'network_info': network_info
            }
        except Exception as e:
            return {
                'type': 'error',
                'message': f'Error procesando JSON: {str(e)}'
            }
    
    def _extract_network_info(self, text):
        """Extrae información de redes del texto"""
        network_info = {
            'subnets': [],
            'ip_ranges': [],
            'routing_info': []
        }
        
        # Patrones para detectar subnets y IPs
        subnet_patterns = [
            r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}/\d{1,2})',  # CIDR notation
            r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})',  # IP address
            r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\s*-\s*\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'  # IP range
        ]
        
        # Buscar patrones de red
        for pattern in subnet_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if '/' in match:  # CIDR
                    network_info['subnets'].append(match)
                elif '-' in match:  # IP range
                    network_info['ip_ranges'].append(match)
                else:  # Single IP
                    network_info['subnets'].append(match)
        
        # Buscar información de ruteo
        routing_patterns = [
            r'(?:ruta|route|gateway|next.?hop|destino|destination)\s*[:=]\s*(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})',
            r'(?:subnet|red|network)\s*[:=]\s*([^\n\r]+)',
            r'(?:máscara|mask|netmask)\s*[:=]\s*(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'
        ]
        
        for pattern in routing_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            network_info['routing_info'].extend(matches)
        
        # Solo retornar si encontramos información de red
        if network_info['subnets'] or network_info['ip_ranges'] or network_info['routing_info']:
            return network_info
        return None
    
    def _detect_network_excel(self, headers, data):
        """Detecta si el Excel contiene información de redes"""
        network_keywords = ['subnet', 'ip', 'red', 'network', 'gateway', 'ruta', 'máscara', 'mask', 'cidr']
        
        # Verificar si los headers contienen palabras clave de red
        header_text = ' '.join(headers).lower()
        if any(keyword in header_text for keyword in network_keywords):
            return {
                'type': 'excel_network',
                'headers': headers,
                'data': data
            }
        return None
    
    def _detect_network_csv(self, headers, data):
        """Detecta si el CSV contiene información de redes"""
        network_keywords = ['subnet', 'ip', 'red', 'network', 'gateway', 'ruta', 'máscara', 'mask', 'cidr']
        
        # Verificar si los headers contienen palabras clave de red
        header_text = ' '.join(headers).lower()
        if any(keyword in header_text for keyword in network_keywords):
            return {
                'type': 'csv_network',
                'headers': headers,
                'data': data
            }
        return None
    
    def _detect_network_json(self, json_data):
        """Detecta si el JSON contiene información de redes"""
        def search_network_data(obj, path=""):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    current_path = f"{path}.{key}" if path else key
                    if isinstance(value, str):
                        # Buscar patrones de IP/subnet
                        if re.search(r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', value):
                            return True
                    elif isinstance(value, (dict, list)):
                        if search_network_data(value, current_path):
                            return True
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    current_path = f"{path}[{i}]"
                    if search_network_data(item, current_path):
                        return True
            return False
        
        if search_network_data(json_data):
            return {
                'type': 'json_network',
                'data': json_data
            }
        return None
